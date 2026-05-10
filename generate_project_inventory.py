#!/usr/bin/env python3
"""
Walk the project tree and write:
  ATLAS_PROJECT_FILE_INVENTORY.md   — paste into Gemini / docs
  ATLAS_PROJECT_FILE_INVENTORY.docx — Word version

Excludes: .git, __pycache__, .cursor, node_modules (if present)
"""
from __future__ import annotations

import re
from datetime import datetime, timezone
from pathlib import Path

# Optional: pip install python-docx
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

SKIP_DIR_NAMES = {".git", "__pycache__", ".cursor", "node_modules", ".venv", "venv"}

# Path-relative or filename-specific notes (override generic extension blurb)
KNOWN: dict[str, str] = {
    "api_server.py": "FastAPI app: POST /omega, /query, /research, positions, watchlist, GET /health, /v2 dashboard (v4-first).",
    "atlas_omega.py": "OmegaAgent — cross-domain financial agent: classifier, parallel data workers, stock_universe discovery, Gemini JSON synthesis.",
    "query_router.py": "QueryRouter + FourLoopEngine — 10-loop equity/options pipeline (scrape → synthesize → personalize → rules → scenarios → memory → adversarial → narrative).",
    "web_scraper.py": "Omnivore scraper: Finviz, news, SEC, many sources; builds context blobs for research.",
    "market_scanner.py": "Market regime, squeeze-style metrics, full_awareness for tickers.",
    "stock_universe.py": "Progressive Finviz + yfinance funnel (universe → filter → signals → optional deep rank); omega_discovery_to_scan_params.",
    "deep_research.py": "Heavy ticker / discovery pipeline: scrape + structured Gemini reports; CLI and /research fallback.",
    "delta_reporter.py": "Delta-style HTML reports under reports/.",
    "multi_ranker.py": "Scores/ranks candidates (used by stock_universe pass 4).",
    "gemini_limiter.py": "Global rate limit / spacing for all Gemini calls.",
    "memory.py": "Persistent ticker memory (SQLite); used in research context.",
    "tracker.py": "Recommendation history & outcomes; personalization loops 5 & 8.",
    "rag_engine.py": "SEC/document RAG helpers; Chroma-backed when enabled.",
    "volume_profile.py": "Volume profile POC/VAH/VAL for synthesis context.",
    "sector_tracker.py": "Sector rotation / wind context.",
    "congress_tracker.py": "Congressional trade disclosures; cache under congress_cache/.",
    "options_simulator.py": "Options P/L and structure math.",
    "paper_trader.py": "Paper trading monitor and logging.",
    "broker_alpaca.py": "Alpaca API integration (paper/real).",
    "broker_tradier.py": "Tradier API integration.",
    "position_sizer.py": "Position sizing utilities.",
    "alerts.py": "Price/user alerts; notifications (winsound guarded).",
    "auto_bot.py": "Scheduled / intraday bot, scans, LIVE_DASHBOARD generation.",
    "auto_tuner.py": "Config/weight tuning surface for agents.",
    "dashboard_server.py": "HTTP 8765: dashboard.html, /state, refresh, reports, /v2 financial UI.",
    "dashboard.html": "Legacy live portfolio / regime HTML dashboard.",
    "news_scanner.py": "News scanning loop / CLI.",
    "screen_watcher.py": "Screen capture watcher pipeline.",
    "playwright_scraper.py": "JS-heavy sites via Playwright.",
    "backtest_sandbox.py": "Backtesting experiments.",
    "self_coder.py": "Self-modification / codegen experiments.",
    "START_ATLAS.py": "Launches auto_bot --watch + dashboard_server; opens browser.",
    "test_omega.py": "Integration test: /query or /omega + health; saves test_result_soun.json.",
    "build_audit_docx.py": "Converts ATLAS_FULL_AGENT_AUDIT.md → .docx.",
    "generate_project_inventory.py": "This script — regenerates project file inventory.",
    "requirements.txt": "Python dependencies (FastAPI, yfinance, genai, playwright, etc.).",
    ".gitignore": "Git ignore rules.",
    ".env": "Secrets — GOOGLE_API_KEY, broker keys (never commit).",
    "positions_cache.json": "Manual stock/option positions for API and loop 5.",
    "watchlist.json": "Watchlist tickers for API.",
    "paper_trades.json": "Paper trade log for personalization.",
    "dashboard_state.json": "Cached JSON for dashboard_server UI.",
    "atlas_alerts.json": "Alert definitions persisted.",
    "atlas_alerts.log": "Alert log (may be empty).",
    "atlas_tracking_state.json": "Auto-tracker / bot tracking state.",
    "atlas_pending_deep.json": "Queue or pending deep research jobs.",
    "research_history.json": "History of research runs.",
    "weekly_insight.json": "Weekly insight artifact for dashboards/bot.",
    "test_result_soun.json": "Saved API test response (e.g. Omega or /query).",
    "atlas_memory.db": "SQLite — long-lived memories.",
    "atlas_tracker.db": "SQLite — recommendations / P&L history.",
    "atlas_dashboard_v2.html": "Financial Intelligence UI (older; light theme).",
    "atlas_dashboard_v3.html": "Alternative dark / bento-style UI experiment.",
    "atlas_dashboard_v4.html": "Primary SPA dashboard — dark/light toggle; calls /omega then /query.",
    "ATLAS_OUTPUT_MAP.html": "Static map / overview of outputs.",
    "ATLAS_OVERVIEW.html": "Static overview HTML.",
    "LIVE_DASHBOARD.html": "Generated live tactical dashboard (auto_bot).",
    "LIVE_REPORTS.html": "Index of live reports.",
    "CANVAS_1_Roadmap.html": "Roadmap canvas / presentation HTML.",
    "CANVAS_3_ATLAS_vs_World.html": "Comparison canvas HTML.",
    "CANVAS_4_Intelligence_Layers.html": "Architecture layers canvas HTML.",
    "CLAUDE.md": "Long internal spec / task list for ATLAS development.",
    "Prompt.md": "Short prompt notes (e.g. /omega vs /query testing).",
    "ATLAS_FULL_AGENT_AUDIT.md": "Technical + product audit of the agent (Markdown).",
    "ATLAS_FULL_AGENT_AUDIT.docx": "Word export of the audit.",
    "ATLAS_v3_Full_Audit.docx": "Earlier audit / notes (Word).",
    "ATLAS_Master_Roadmap_to_1B (1).docx": "Business roadmap document.",
    "chroma.sqlite3": "Chroma vector DB (RAG embeddings metadata + segments).",
    "data_level0.bin": "Chroma HNSW / vector storage binary.",
    "header.bin": "Chroma index header.",
    "length.bin": "Chroma index lengths.",
    "link_lists.bin": "Chroma graph link storage.",
    "backtest_cache.json": "RAG folder cache for backtests.",
    "ingested.json": "RAG ingestion manifest (which docs embedded).",
}


def ext_blurb(suffix: str) -> str:
    s = suffix.lower()
    return {
        ".py": "Python source module or script.",
        ".md": "Markdown documentation.",
        ".txt": "Plain text.",
        ".html": "HTML — static page or generated report.",
        ".htm": "HTML report.",
        ".json": "JSON data — config, cache, or API state.",
        ".db": "SQLite database file.",
        ".sqlite3": "SQLite database (Chroma).",
        ".bin": "Binary index / embedding chunk data (Chroma).",
        ".log": "Log file.",
        ".jpg": "JPEG image (e.g. watcher capture).",
        ".png": "PNG image asset.",
        ".docx": "Microsoft Word document.",
        ".env": "Environment variables (secrets).",
    }.get(s, f"File type {s or '.'} — see project context.")


def should_skip_dir(path: Path) -> bool:
    return path.name in SKIP_DIR_NAMES


def describe(rel: Path, size: int) -> str:
    key = rel.name
    if key in KNOWN:
        return KNOWN[key]
    # path-specific overrides
    rs = rel.as_posix()
    if rs.startswith("reports/"):
        return "Generated equity research / delta HTML report."
    if rs.startswith("deep_reports/"):
        if rs.endswith("_research_log.json"):
            return "JSON log for a deep research run."
        return "Saved deep-research HTML output."
    if rs.startswith("delta_snapshots/"):
        return "JSON snapshot of delta / ticker state."
    if rs.startswith("congress_cache/"):
        return "Cached congressional trades JSON."
    if "atlas_memory_data" in rs and rs.endswith(".jpg"):
        return "Screen / watcher capture image."
    if "atlas_rag" in rs and rel.suffix == ".bin":
        return "Chroma vector index binary (do not edit manually)."
    return ext_blurb(rel.suffix)


def collect_files(root: Path) -> list[tuple[Path, int]]:
    out: list[tuple[Path, int]] = []
    for p in root.rglob("*"):
        if p.is_dir():
            if should_skip_dir(p):
                # prune: don't descend — but rglob already mixed; filter files only
                continue
        if not p.is_file():
            continue
        try:
            if any(part in SKIP_DIR_NAMES for part in p.relative_to(root).parts):
                continue
        except ValueError:
            continue
        try:
            sz = p.stat().st_size
        except OSError:
            sz = -1
        out.append((p, sz))
    out.sort(key=lambda x: str(x[0].relative_to(root)).lower())
    return out


def write_markdown(root: Path, rows: list[tuple[Path, int]], md_path: Path) -> None:
    lines: list[str] = [
        "# ATLAS Project — Complete File Inventory (for LLM context)",
        "",
        f"**Generated:** {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}  ",
        "**Purpose:** Single document listing **every tracked file** in the repo (excluding `.git`, `__pycache__`, `.cursor`, `venv`) with a short purpose. Use when the full codebase is too large to paste.",
        "",
        "**Security:** Listing mentions `.env` — **never paste secrets** into an LLM; redact or omit that file when sharing.",
        "",
        "**How to regenerate:** `python generate_project_inventory.py`",
        "",
        "---",
        "",
        "## Summary counts",
        "",
        f"- **Total files:** {len(rows)}",
        "",
        "---",
        "",
        "## File listing (path → size → description)",
        "",
    ]
    for p, sz in rows:
        rel = p.relative_to(root)
        rs = rel.as_posix()
        szs = f"{sz:,} bytes" if sz >= 0 else "unknown"
        lines.append(f"### `{rs}`")
        lines.append(f"- **Size:** {szs}")
        lines.append(f"- **Role:** {describe(rel, sz)}")
        lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("*End of inventory.*")
    md_path.write_text("\n".join(lines), encoding="utf-8")


def write_docx(root: Path, rows: list[tuple[Path, int]], docx_path: Path) -> None:
    if not HAS_DOCX:
        return
    doc = Document()
    t = doc.add_heading("ATLAS Project — File Inventory", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph(
        f"Generated {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')} — "
        "for Gemini / stakeholders when the repo is too large to load whole."
    )
    doc.add_paragraph(
        "Security: This inventory lists .env — never paste API keys or broker secrets into an LLM."
    )
    doc.add_paragraph("Regenerate: python generate_project_inventory.py")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Relative path"
    hdr[1].text = "Size (bytes)"
    hdr[2].text = "Description"

    for p, sz in rows:
        rel = p.relative_to(root)
        row = table.add_row().cells
        row[0].text = rel.as_posix()
        row[1].text = str(sz) if sz >= 0 else "?"
        row[2].text = describe(rel, sz)

    doc.save(docx_path)


def main() -> None:
    root = Path(__file__).resolve().parent
    rows = collect_files(root)
    md_path = root / "ATLAS_PROJECT_FILE_INVENTORY.md"
    docx_path = root / "ATLAS_PROJECT_FILE_INVENTORY.docx"

    write_markdown(root, rows, md_path)
    print(f"Wrote {md_path} ({len(rows)} files)")

    if HAS_DOCX:
        write_docx(root, rows, docx_path)
        print(f"Wrote {docx_path}")
    else:
        print("Skip .docx — install: pip install python-docx")


if __name__ == "__main__":
    main()
