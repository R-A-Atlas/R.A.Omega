#!/usr/bin/env python3
"""One-off: ATLAS_FULL_AGENT_AUDIT.md -> ATLAS_FULL_AGENT_AUDIT.docx"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_runs_with_formatting(paragraph, text: str) -> None:
    """Split on **bold** and `code` (simple linear pass)."""
    if not text:
        return
    # Merge ** and ` into ordered tokens
    parts: list[tuple[str, str]] = []  # (kind, content) kind in normal|bold|code
    i = 0
    while i < len(text):
        if text.startswith("**", i):
            j = text.find("**", i + 2)
            if j != -1:
                parts.append(("bold", text[i + 2 : j]))
                i = j + 2
                continue
        if text[i] == "`" and i + 1 < len(text):
            j = text.find("`", i + 1)
            if j != -1:
                parts.append(("code", text[i + 1 : j]))
                i = j + 1
                continue
        # collect until next special
        next_b = text.find("**", i)
        next_c = text.find("`", i)
        candidates = [p for p in (next_b, next_c) if p != -1]
        end = min(candidates) if candidates else len(text)
        if end > i:
            parts.append(("normal", text[i:end]))
        i = end if end > i else i + 1

    for kind, chunk in parts:
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        if kind == "bold":
            run.bold = True
        if kind == "code":
            run.font.name = "Consolas"
            run.font.size = Pt(9)


def is_table_separator(line: str) -> bool:
    s = line.strip()
    if "|" not in s:
        return False
    return bool(re.match(r"^\|?[\s\-:|]+\|?$", s))


def parse_table_row(line: str) -> list[str]:
    row = [c.strip() for c in line.strip().strip("|").split("|")]
    return row


def main() -> None:
    base = Path(__file__).resolve().parent
    md_path = base / "ATLAS_FULL_AGENT_AUDIT.md"
    out_path = base / "ATLAS_FULL_AGENT_AUDIT.docx"
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    in_code = False
    code_buf: list[str] = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Consolas"
                run.font.size = Pt(8)
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if line.startswith("# ") and not line.startswith("##"):
            h = doc.add_heading(line[2:].strip(), level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            i += 1
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue

        if line.strip() == "---":
            i += 1
            continue

        # Table: header row, separator, body
        if (
            "|" in line
            and i + 1 < len(lines)
            and is_table_separator(lines[i + 1])
        ):
            header = parse_table_row(line)
            i += 2
            rows: list[list[str]] = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                if is_table_separator(lines[i]):
                    i += 1
                    continue
                rows.append(parse_table_row(lines[i]))
                i += 1
            tbl = doc.add_table(rows=len(rows) + 1, cols=len(header))
            tbl.style = "Table Grid"
            for c, h in enumerate(header):
                tbl.rows[0].cells[c].text = re.sub(r"\*+", "", h).strip()
            for r, row in enumerate(rows):
                for c, cell in enumerate(row):
                    if c < len(tbl.rows[r + 1].cells):
                        tbl.rows[r + 1].cells[c].text = re.sub(r"\*+", "", cell).strip()
            doc.add_paragraph()
            continue

        if line.strip().startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs_with_formatting(p, line.strip()[2:])
            i += 1
            continue

        m = re.match(r"^(\d+)\.\s+(.+)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_runs_with_formatting(p, m.group(2))
            i += 1
            continue

        if line.strip():
            p = doc.add_paragraph()
            add_runs_with_formatting(p, line.strip())
        i += 1

    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
